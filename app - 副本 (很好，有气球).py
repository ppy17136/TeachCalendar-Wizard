import os
import streamlit as st
import pdfplumber
import fitz  # PyMuPDF
from docx import Document
import mammoth
import requests
import re
import numpy as np
import matplotlib.pyplot as plt
from openai import OpenAI
import base64
import io
from PIL import Image
import google.generativeai as genai
import json
from docxtpl import DocxTemplate  # 必须安装 docxtpl
from datetime import datetime
# 签名插入示例
from docxtpl import InlineImage
from docx.shared import Mm, Pt

# --- 1. 基础环境与配置 ---
plt.rcParams['font.family'] = ['SimHei', 'sans-serif']
plt.rcParams['axes.unicode_minus'] = False

st.set_page_config(page_title="智能教学辅助系统", layout="wide", initial_sidebar_state="expanded")

# --- 3. 密钥获取与侧边栏 ---
BACKEND_QWEN_KEY = st.secrets.get("QWEN_API_KEY", "")
BACKEND_GEMINI_KEY = st.secrets.get("GEMINI_API_KEY", "")

# --- 2. 状态自动化初始化 (防止变量未定义报错) ---
# 初始化全局会话状态
if "score_records" not in st.session_state:
    st.session_state.score_records = []
if "generated_syllabus" not in st.session_state:
    st.session_state.generated_syllabus = None
if "generated_calendar" not in st.session_state:
    st.session_state.generated_calendar = None
if "generated_program" not in st.session_state:
    st.session_state.generated_program = None
# 使用 setdefault 确保变量一定存在
st.session_state.setdefault("score_records", [])
st.session_state.setdefault("gen_content", {"syllabus": None, "calendar": None, "program": None})
# --- 3. 侧边栏：引擎切换与密钥管理 ---
with st.sidebar:
    st.header("⚙️ 模型引擎设置")
    selected_provider = st.radio("选择主 AI 引擎", ["Gemini", "Qwen (通义千问)"])
    
    ACTIVE_QWEN_KEY = BACKEND_QWEN_KEY
    ACTIVE_GEMINI_KEY = BACKEND_GEMINI_KEY

    if selected_provider == "Gemini":
        user_gem_key = st.text_input("填写 Gemini API Key (可选)", type="password", help="留空则使用后台默认 Key")
        if user_gem_key: ACTIVE_GEMINI_KEY = user_gem_key
        selected_model = st.selectbox("版本", ["gemini-2.5-flash", "gemini-2.0-flash-exp", "gemini-2.5-pro"])
        engine_id = "Gemini"
        if ACTIVE_GEMINI_KEY: 
            genai.configure(api_key=ACTIVE_GEMINI_KEY)
        else:
            st.error("⚠️ 未检测到有效 Gemini Key")
    else:
        user_qw_key = st.text_input("填写 Qwen API Key (可选)", type="password", help="留空则使用后台默认 Key")
        if user_qw_key: ACTIVE_QWEN_KEY = user_qw_key
        selected_model = st.selectbox("版本", ["qwen-plus", "qwen-max", "qwen-turbo"])
        engine_id = "Qwen"
        if not ACTIVE_QWEN_KEY:
            st.error("⚠️ 未检测到有效 Qwen Key")

    st.divider()
    st.info(f"💡 当前模式：使用 **{engine_id}** 处理。")
    # 侧边栏底部也可以加提示
    st.caption("🖥️ 建议环境：Google Chrome 浏览器")
    
# --- 4. 核心功能函数 --- 
def create_docx(text):
    """将文本转换为可下载的 Word，彻底清洗所有标记"""
    doc = Document()
    
    # 1. 首先通过正则表达式清除所有 HTML 标签 (如 <br/>)
    # 2. 接着通过链式 replace 清除 Markdown 的标题号和加粗符号
    clean_text = re.sub('<[^<]+?>', '', text) \
                   .replace("### ", "") \
                   .replace("## ", "") \
                   .replace("# ", "") \
                   .replace("**", "")
    
    # 写入 Word
    for line in clean_text.split('\n'):
        if line.strip(): # 过滤掉多余的空行
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(12)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def ai_generate(prompt, provider, model_name):
    """统一文本生成接口"""
    if provider == "Gemini":
        if not ACTIVE_GEMINI_KEY: return "错误：未配置密钥"
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            return response.text
        except Exception as e: return f"Gemini 失败: {str(e)}"
    else:
        if not ACTIVE_QWEN_KEY: return "错误：未配置密钥"
        client = OpenAI(api_key=ACTIVE_QWEN_KEY, base_url="https://dashscope.aliyuncs.com/compatible-mode/v1")
        try:
            completion = client.chat.completions.create(model=model_name, messages=[{"role": "user", "content": prompt}])
            return completion.choices[0].message.content
        except Exception as e: return f"Qwen 失败: {str(e)}"

def ai_ocr(image_bytes, provider, model_name):
    """根据引擎进行图片文字识别"""
    if provider == "Gemini":
        if not ACTIVE_GEMINI_KEY: return "错误：未配置密钥"
        try:
            model = genai.GenerativeModel(model_name)
            res = model.generate_content(["识别并输出图中文字内容。若是试卷，请提取题目和回答。", {"mime_type": "image/jpeg", "data": image_bytes}])
            return res.text
        except Exception as e: return f"Gemini 视觉识别失败: {str(e)}"
    else:
        if not ACTIVE_QWEN_KEY: return "错误：未配置密钥"
        # 图片压缩优化
        img = Image.open(io.BytesIO(image_bytes)).convert('RGB')
        max_width = 1024
        if img.width > max_width:
            scale = max_width / img.width
            img = img.resize((max_width, int(img.height * scale)))
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=90)
        b64img = base64.b64encode(buf.getvalue()).decode("utf-8")
        
        client = OpenAI(api_key=ACTIVE_QWEN_KEY, base_url="https://dashscope.aliyuncs.com/compatible-mode/v1")
        try:
            completion = client.chat.completions.create(
                model="qwen-vl-ocr-latest",
                messages=[{"role": "user", "content": [{"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64img}"}}, {"type": "text", "text": "请提取图中所有文字内容"}]}]
            )
            return completion.choices[0].message.content
        except Exception as e: return f"Qwen OCR 失败: {str(e)}"

# --- 5. 文档与工具函数 ---
def extract_text_from_file(file):
    """支持多格式文本提取"""
    try:
        if file.name.endswith(".docx"):
            return "\n".join([p.text for p in Document(file).paragraphs])
        elif file.name.endswith(".pdf"):
            with pdfplumber.open(file) as pdf:
                return "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif file.name.endswith(".doc"):
            return mammoth.convert_to_text(file).value
        return "格式暂不支持"
    except Exception as e:
        return f"解析失败: {str(e)}"


def safe_extract_text(file, max_chars=15000):
    if not file: return ""
    try:
        text_list = []
        if file.name.endswith(".pdf"):
            with fitz.open(stream=file.read(), filetype="pdf") as doc:
                for page in doc:
                    text_list.append(page.get_text())
                    if sum(len(t) for t in text_list) > max_chars: break
            return "".join(text_list)[:max_chars]
            
        elif file.name.endswith(".docx"):
            doc = Document(file)
            for p in doc.paragraphs:
                if p.text.strip(): text_list.append(p.text)
            
            for table in doc.tables:
                for row in table.rows:
                    processed_cells = []
                    for cell in row.cells:
                        content = cell.text
                        # --- 核心改进：非互斥全量替换，涵盖更多 Word 特殊符号 ---
                        # 识别“已选中”符号
                        checked_chars = ['☑', 'þ', '\xfe', '\uf0fe', '☒', '√']
                        # 识别“未选中”符号
                        unchecked_chars = ['☐', '¨', '\xa8', '\uf0a1', '□']
                        
                        for c in checked_chars:
                            content = content.replace(c, '[已选中]')
                        for u in unchecked_chars:
                            content = content.replace(u, '[未选中]')
                        
                        processed_cells.append(content.strip())
                    
                    row_text = [c for c in processed_cells if c]
                    if row_text: text_list.append(" | ".join(row_text))
            
            return "\n".join(text_list)[:max_chars]
        elif file.name.endswith(".doc"):
            return mammoth.convert_to_text(file).value[:max_chars]            
        return ""

    except Exception as e:
        st.error(f"文件 {file.name} 解析出错: {str(e)}")
        return ""

def render_pdf_images(pdf_file):
    images = []
    pdf_file.seek(0)
    with fitz.open(stream=pdf_file.read(), filetype="pdf") as pdf:
        for page in pdf:
            pix = page.get_pixmap(matrix=fitz.Matrix(2,2))
            images.append(pix.tobytes("png"))
    return images

def nav_bar(show_back=False):
    st.markdown(f'<div style="background:#1E2129;padding:20px;border-radius:10px;margin-bottom:10px;"><h1 style="color:white;margin:0;font-size:24px;">🎓 智能教学与批卷系统 <span style="font-size:14px;color:#888;">{engine_id} 引擎在线</span></h1></div>', unsafe_allow_html=True)
    if show_back:
        if st.button("⬅️ 返回主页", use_container_width=True):
            st.query_params["page"] = "首页"
            st.rerun()

# --- 6. 页面功能定义 ---
def page_home():
    nav_bar()
    st.markdown("### 🛠️ 教务与批改功能矩阵")
    cols = st.columns(3)
    modules = [
        ("📄", "教学大纲生成", "大纲"), ("📅", "教学日历生成", "日历"), ("📋", "培养方案生成", "方案"),
        ("📝", "智能批卷系统", "批卷"), ("📈", "成绩分析报告", "分析"), ("⚙️", "系统设置", "设置")
    ]
    for i, (icon, title, link) in enumerate(modules):
        with cols[i % 3]:
            st.markdown(f'<div style="border:1px solid #ddd;padding:20px;border-radius:10px;text-align:center;"><span style="font-size:40px;">{icon}</span><h4>{title}</h4></div>', unsafe_allow_html=True)
            if st.button(f"进入{title}", key=f"nav_{i}", use_container_width=True):
                st.query_params["page"] = link
                st.rerun()

def page_syllabus():
    nav_bar(show_back=True)
    st.subheader("📄 深度智造：教学大纲 (支持上传教材分析)")
    
    # 5.1 上传辅助资料区域
    with st.expander("##### 📚 第一步：上传参考资料 (教材/培养方案/参考文献)", expanded=True):
        col_u1, col_u2 = st.columns(2)
        book_file = col_u1.file_uploader("上传教材/参考书 PDF/Word", type=["pdf", "docx"])
        plan_file = col_u2.file_uploader("上传人才培养方案 PDF/Word", type=["pdf", "docx"])
        
    # 5.2 手工填写基本信息
    with st.form("syllabus_form"):
        st.markdown("##### 📚 第二步：填写关键参数")        
        # 第一排：基础课程信息 
        c1, c2, c3 = st.columns(3)
        name = c1.text_input("课程名称", value="数值模拟在材料成型中的应用")
        major = c2.text_input("适用专业", value="材料成型及控制工程（焊接方向）")
        course_type = c3.selectbox("课程性质", ["必修", "限选", "选修"], index=1)

        # 第二排：学分学时与考核 
        c4, c5, c6 = st.columns(3)
        hours = c4.number_input("总学时", value=32)
        credits = c5.number_input("总学分", value=2.0, step=0.5)
        assessment = c6.selectbox("考核方式", ["考试", "考查"], index=1)

        # 第三排：学期与要求 
        c7, c8 = st.columns(2)
        semester = c7.selectbox("开课学期", ["一", "二", "三", "四", "五", "六", "七", "八"], index=4)
        prerequisites = c8.text_area("先修课程要求", value="高等数学、工程力学，具备基本微积分和工程力学知识", height=68)

        # 核心目标与思政
        obj = st.text_area("培养目标", placeholder="输入课程培养目标...", value="课程目标1：能够了解材料成型的数值模拟软件的原理和方法，并理解其局限性；\n课程目标2：能够选用合适的专业数值模拟软件分析材料成型工程中的复杂问题；\n课程目标3：能够选用适合的数值模拟软件预测材料成型工程问题，并分析其局限性。")
        ideology = st.text_area("思政融入点", value="国产工业软件发展、两弹一星精神")

        if st.form_submit_button("🚀 结合上传资料生成 OBE 标准大纲"):
            with st.spinner("正在阅读文档并构思大纲..."):
                #book_ctx = extract_text_from_file(book_file) if book_file else "未提供教材"
                plan_ctx = extract_text_from_file(plan_file) if plan_file else "未提供培养方案"   
                book_ctx = safe_extract_text(book_file) if book_file else "未提供教材"
                #plan_ctx = safe_extract_text(plan_file) if plan_file else "未提供培养方案"
                
                prompt = f"""
                        你是一位资深的高校工程教育认证专家。请为《{name}》课程撰写一份高质量教学大纲。文字专业且符合OBE理念。
                        
                        **严格排版要求：**
                        1. 禁止使用任何 HTML 标签（如 <br/>, <b> 等）。
                        2. 所有的表格必须使用标准 Markdown 格式：| 列1 | 列2 |。
                        3. 必须包含分隔线：| :--- | :--- |。
                        4. 每个表格上方和下方必须各留一行空行。
                        
                        **背景资料（请务必参考以下内容）：**
                        1. 教材/内容核心：{book_ctx[:12000]} (注：由于长度限制，已截取前1万字符)
                        2. 专业培养要求：{plan_ctx[:10000]}
                        
                        **手工填写的参数：**                    
                        - 课程性质：{course_type} | 考核方式：{assessment} | 学分：{credits} | 学时：{hours}
                        - 适用专业：{major} | 思政：{ideology} | 开课学期{semester} | 先修课程及其要求{prerequisites}                   
                        - 课程目标支撑毕业要求表（含课程目标{obj}
                        
                        **大纲必须包含：**
                        - 课程基本信息表，包含大纲名称、课程名称{name}、英文名称、编码、课程性质{course_type}、适用专业{major}、考核方式{assessment}、总学分{credits}、总 学 时{hours}（理论学时X、实验学时X、实训学时X、其他（讨论）	学时X）、开课学期{semester}、先修课程及其要求{prerequisites}等
                        - 课程简介（理实结合，不少于200字）
                        - 建议教材	 
                        - 参考资料	 
                        - 教学条件
                        - 课程目标支撑毕业要求表（含课程目标{obj}、支撑指标点如4.1/5.1及支撑强度H/M/L）
                        - 德育目标
                        - 教学内容学时分配表（确保总学时为{hours}）（教学内容参考教材和参考材料{book_ctx}，包含序号、教学内容、学生学习预期成果、计划学时、支撑目标、教学方式、其它（作业、习题、实验等）
                        - 课程目标考核
                        - 课程目标达成情况评价
                        - 考核评价表（包含平时成绩与期末考试占比）                    
                        - 课程考核，包含标准考试评分标准、作业评分标准
                        - 大作业评分标准，包含作业内容、评价标准（90-100分	70-89 分	60-69分	0-59分）、所占比重
                        - 课程思政实施方案（结合：{ideology}），包含思政内容切入点、典型案例、教育载体及方法、预期达到的目标、	体现的价值观或思政元素
                        
                        **尤其注意构建《课程目标支撑毕业要求表》时：**
                        请基于培养方案{plan_ctx}严格以下对应关系生成表格，禁止随意发挥：
                        1. 课程目标1：{obj.split('课程目标2')[0] if '课程目标2' in obj else obj} 
                           --> 必须支撑：5.1 (工具使用)。
                        2. 课程目标2：... (以此类推，请解析用户输入的 {obj})

                        **表格格式要求：**
                        | 课程目标 | 支撑毕业要求及指标点 | 支撑强度 (H/M/L) |
                        | :--- | :--- | :--- |
                        | 课程目标1：[简述目标内容] | 5.1 了解常用现代仪器... | H |
                        | 课程目标2：[简述目标内容] | 5.2 能够选择与使用恰当仪器... | M |

                        **特别注意：**
                        - 每一行只能对应一个课程目标。
                        - 每一个课程目标只能对应一个毕业要求及指标点
                        - 指标点描述必须完整。
                        - 支撑强度必须根据该目标对指标点的支撑力度给出唯一的 H、M 或 L。                        
                        """            
                # 执行生成并存入缓存
                st.session_state.gen_content["syllabus"] = ai_generate(prompt, engine_id, selected_model)
                st.session_state['course_name'] = name
                st.session_state['total_hours'] = hours
                st.session_state['major'] = major # 适用专业
                #st.session_state['assessment_method'] = assessment # 考核方式
                st.session_state['course_objectives'] = obj # 存储原始输入的课程目标文本
                st.session_state['ideology_points'] = ideology # 存储思政点，以便日历中安排思政课次                

                st.success("✅ 大纲生成成功！")

    if st.session_state.gen_content["syllabus"]:
        st.markdown("---")
        st.container(border=True).markdown(st.session_state.gen_content["syllabus"])
        col1, col2 = st.columns(2)
        col1.download_button("💾 下载 Word 版大纲", create_docx(st.session_state.gen_content["syllabus"]), file_name=f"{name}_大纲.docx")
        col2.download_button("📝 下载文本版 (TXT)", st.session_state.gen_content["syllabus"], file_name=f"{name}_大纲.txt")        



# ==================== 1. 核心渲染与辅助函数 ====================
# --- 辅助函数：读取模版结构 ---
def read_local_docx_structure(file_path):
    if not os.path.exists(file_path):
        return "模版文件不存在"
    try:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if "{{" in p.text])
    except:
        return "模版读取失败"

# --- 核心函数：渲染 Word 文档 ---
def render_calendar_docx(template_path, data_dict, sig_images=None):
    """
    data_dict: 包含所有标签键值的字典
    sig_images: 字典，格式为 {"标签名": 文件流}
    """
    try:
        doc = DocxTemplate(template_path)
        
        # 1. 递归清洗数据中的 None 或 N/A
        def clean_val(v):
            if v is None or str(v).lower() in ["none", "n/a", "未提供"]: return ""
            return v

        processed_data = {}
        for k, v in data_dict.items():
            if k == "schedule": # 进度表特殊处理
                processed_data[k] = [{sk: clean_val(sv) for sk, sv in item.items()} for item in v]
            else:
                processed_data[k] = clean_val(v)

        # 2. 注入签名图片
        if sig_images:
            for key, img_stream in sig_images.items():
                if img_stream:
                    # 将上传的图片转换为 Word 内部对象，宽度设为 30mm
                    processed_data[key] = InlineImage(doc, img_stream, width=Mm(30))
                else:
                    processed_data[key] = ""

        # 3. 渲染并导出
        doc.render(processed_data, autoescape=True)
        target_stream = io.BytesIO()
        doc.save(target_stream)
        return target_stream.getvalue()
    except Exception as e:
        st.error(f"渲染失败: {str(e)}")
        return None

# --- 教师视图：编报与拆分 ---
def render_teacher_view():
    st.markdown("#### 📝 教师端：日历编报与原文对照")
    
    # 1. 基础信息配置 (全项列出)
    with st.container(border=True):
        st.markdown("##### 👤 1. 基础信息设置")
        c1, c2, c3 = st.columns([1.5, 2, 1.5])
        school_name = c1.text_input("学校名称", value="辽宁石油化工大学")
        course_name = c2.text_input("课程名称", value=st.session_state.get('course_name', ""))
        
        # 自动日期
        today_str = datetime.now().strftime("%Y年 %m月 %d日")
        
        t1, t2, t3 = st.columns(3)
        teacher_name = t1.text_input("主讲教师姓名", value=st.session_state.get('teacher_name', ""))
        teacher_title = t2.text_input("职称", value=st.session_state.get('teacher_title', ""))
        st.info(f"📅 签名日期：{today_str}")

        h1, h2, h3 = st.columns(3)
        total_hours = h1.number_input("总学时", value=int(st.session_state.get('total_hours', 24)))
        total_weeks = h2.number_input("总周数", value=12)
        current_assessment = h3.radio("考核方式", ["考试", "考查"], horizontal=True, 
                                     index=1 if st.session_state.get('assessment_method') == "考查" else 0)

    # 2. 签名管理
    with st.expander("✍️ 电子签名管理"):
        sig_col1, sig_col2 = st.columns(2)
        use_saved = sig_col1.checkbox("使用预存签名", value=True)
        teacher_sig_file = None
        if not use_saved:
            teacher_sig_file = sig_col2.file_uploader("上传手写签名图片", type=['png', 'jpg'], key="t_sig")
        else:
            st.caption("已关联默认签名：teacher_sig.png")

    # 3. 智能抽取与学时拆分
    st.divider()
    st.markdown("##### 🗓️ 2. 进度表编辑 (支持 OBE 逻辑拆分)")
    syllabus_file = st.file_uploader("上传新教学大纲 (可选)", type=['docx', 'pdf'])
    
    if st.button("🪄 依据大纲抽取进度 (学时 > 2 自动拆分)"):
        with st.spinner("正在对齐大纲与课次..."):
            syl_ctx = safe_extract_text(syllabus_file) if syllabus_file else st.session_state.gen_content.get("syllabus", "")
            prompt = f"""
            解析大纲内容：{syl_ctx[:8000]}
            生成 JSON 列表。要求：
            1. 若某模块学时 > 2，必须拆分为连续行。如 4 学时拆为“模块X(1/2)”2学时和“模块X(2/2)”2学时。
            2. 必须包含 source_text 字段，存入该项对应的大纲原文。
            3. JSON 键：week, sess, content, req, hrs, method, other, obj, source_text
            """
            res = ai_generate(prompt, engine_id, selected_model)
            try:
                st.session_state.calendar_data = json.loads(re.search(r'\[.*\]', res, re.DOTALL).group(0))
            except: st.error("AI 解析失败，请手动填写")

    # 4. 可编辑表格
    if "calendar_data" in st.session_state:
        st.session_state.calendar_data = st.data_editor(
            st.session_state.calendar_data,
            column_config={
                "source_text": st.column_config.TextColumn("📖 大纲原文依据", width="medium", help="此项内容抽取的原始文本"),
                "content": st.column_config.TextColumn("教学内容", width="large"),
                "hrs": st.column_config.NumberColumn("学时", min_value=1, max_value=4)
            },
            num_rows="dynamic", use_container_width=True
        )

    # 5. 提交审批
    if st.button("📤 提交教学日历审批", type="primary", use_container_width=True):
        st.session_state.calendar_status = "Pending_Head"
        st.session_state.calendar_final_data = {
            "school_name": school_name, "course_name": course_name, "teacher_name": teacher_name,
            "teacher_title": teacher_title, "total_hours": total_hours, "total_weeks": total_weeks,
            "assessment_method": current_assessment, "sign_date_1": today_str,
            "schedule": st.session_state.calendar_data
        }
        st.session_state.teacher_sig = teacher_sig_file
        st.success("已提交至系主任审批！")

# --- 审批视图：流转控制 ---
def render_approval_view(role):
    st.markdown(f"#### 🛡️ {'系主任' if role == 'Department_Head' else '主管院长'}审批界面")
    
    target_status = "Pending_Head" if role == "Department_Head" else "Pending_Dean"
    
    if st.session_state.get("calendar_status") == target_status:
        data = st.session_state.calendar_final_data
        st.info(f"待处理申请：{data['course_name']} (教师：{data['teacher_name']})")
        
        with st.expander("🔍 预览日历详细内容"):
            st.table(data['schedule'])
            
        with st.form(f"form_{role}"):
            opinion = st.text_area("审批意见", value="同意，准予执行。" if role == "Dean" else "经审核，符合大纲要求。")
            sig_file = st.file_uploader("签署手写签名图片", type=['png', 'jpg'])
            
            c_a, c_b = st.columns(2)
            if c_a.form_submit_button("✅ 批准通过"):
                if role == "Department_Head":
                    st.session_state.calendar_status = "Pending_Dean"
                    st.session_state.head_opinion = opinion
                    st.session_state.head_sig = sig_file
                    st.session_state.head_date = datetime.now().strftime("%Y年 %m月 %d日")
                else:
                    st.session_state.calendar_status = "Approved"
                    st.session_state.dean_opinion = opinion
                    st.session_state.dean_sig = sig_file
                    st.session_state.dean_date = datetime.now().strftime("%Y年 %m月 %d日")
                st.rerun()
            
            if c_b.form_submit_button("❌ 退回修改"):
                st.session_state.calendar_status = "Draft"
                st.warning("已退回修改")
    else:
        st.write("🍵 暂无待办事项。")

# --- 主页面入口 ---
def page_calendar():
    nav_bar(show_back=True)
    st.subheader("📅 教学日历全流程审批系统")
    
    user_role = st.sidebar.selectbox("切换当前角色", ["授课教师", "系主任", "主管院长"])
    
    if user_role == "授课教师":
        render_teacher_view()
    elif user_role == "系主任":
        render_approval_view("Department_Head")
    else:
        render_approval_view("Dean")

    # --- 底部进度监控与下载 ---
    st.divider()
    status_map = {"Draft": 0, "Pending_Head": 33, "Pending_Dean": 66, "Approved": 100}
    curr_step = status_map.get(st.session_state.get("calendar_status", "Draft"), 0)
    st.progress(curr_step)
    
    # 审批过程显示
    with st.expander("🚥 审批流转状态"):
        st.write(f"**当前状态：** {st.session_state.get('calendar_status')}")
        if "head_opinion" in st.session_state:
            st.write(f"**系主任意见：** {st.session_state.head_opinion} ({st.session_state.head_date})")
        if "dean_opinion" in st.session_state:
            st.write(f"**院领导意见：** {st.session_state.dean_opinion} ({st.session_state.dean_date})")

    # 最终下载
    if st.session_state.get("calendar_status") == "Approved":
        st.balloons()
        # 汇总所有签名和意见
        final_data = st.session_state.calendar_final_data
        final_data.update({
            "head_opinion": st.session_state.head_opinion, "head_date": st.session_state.head_date,
            "dean_opinion": st.session_state.dean_opinion, "dean_date": st.session_state.dean_date
        })
        sig_map = {
            "teacher_sig": st.session_state.teacher_sig,
            "head_sig": st.session_state.head_sig,
            "dean_sig": st.session_state.dean_sig
        }
        
        doc_bytes = render_calendar_docx("template_lnpu.docx", final_data, sig_map)
        if doc_bytes:
            st.download_button("📥 下载完整审批版教学日历 (.docx)", data=doc_bytes, 
                               file_name="已通过_教学日历.docx", use_container_width=True)
        
  
def page_program():
    nav_bar(show_back=True)
    st.subheader("📋 专业人才培养方案生成")
    with st.form("program_form"):
        major = st.text_input("专业名称", value="材料成型及控制工程")
        pos = st.text_area("专业特色", value="服务石油化工行业，聚焦焊接成型与无损检测")
        if st.form_submit_button("生成人才培养方案"):
            prompt = f"撰写{major}专业2024级培养方案。含培养目标、12项毕业要求、特色定位({pos})、核心课程。专业严谨。"
            with st.spinner("正在构建方案..."):
                st.session_state.generated_program = ai_generate(prompt, engine_id, selected_model)

    if st.session_state.generated_program:
        st.markdown("---")
        st.container(border=True).markdown(st.session_state.gen_content["program"])
        st.download_button("💾 下载 Word 版培养方案", create_docx(st.session_state.gen_content["program"]), file_name="培养方案.docx")

def page_grading():
    nav_bar(show_back=True)
    st.subheader("📝 智能试卷批阅与评价")
    c1, c2 = st.columns(2)
    with c1:
        q_file = st.file_uploader("1. 上传试题 (PDF/Word)", type=["pdf", "docx"], key="q")
        q_txt = extract_text_from_file(q_file) if q_file else ""
    with c2:
        s_file = st.file_uploader("2. 上传标准答案 (PDF/Word)", type=["pdf", "docx"], key="s")
        s_txt = extract_text_from_file(s_file) if s_file else ""

    st.divider()
    papers = st.file_uploader("3. 批量上传学生卷纸 (图片/PDF)", type=["jpg", "png", "pdf"], accept_multiple_files=True)

    for idx, paper in enumerate(papers or []):
        with st.container(border=True):
            st.write(f"**学生 {idx+1}:** {paper.name}")
            s_name = st.text_input("姓名", value=f"学生_{idx+1}", key=f"sn_{idx}")
            
            ocr_text = ""
            if paper.type == "application/pdf":
                imgs = render_pdf_images(paper)
                for i, img in enumerate(imgs):
                    st.image(img, width=350)
                    with st.expander("🔍 查看高清大图"): st.image(img, use_container_width=True)
                    with st.spinner("识别中..."): ocr_text += ai_ocr(img, engine_id, selected_model) + "\n"
            else:
                img_data = paper.read()
                st.image(img_data, width=350)
                with st.expander("🔍 查看高清大图"): st.image(img_data, use_container_width=True)
                with st.spinner("识别中..."): ocr_text = ai_ocr(img_data, engine_id, selected_model)
            
            final_ans = st.text_area("识别结果校对", value=ocr_text, key=f"ocr_{idx}", height=150)
            
            if st.button(f"🚀 {engine_id} 自动批改", key=f"go_{idx}"):
                with st.spinner("正在评分..."):
                    p = f"题目：{q_txt}\n答案：{s_txt}\n学生：{final_ans}\n请评分(满分100)并给出批注。格式：\n分数：[数字]\n批注：[解析]"
                    res = ai_generate(p, engine_id, selected_model)
                    st.markdown(res)
                    score = int(re.search(r"分数[：:]\s*(\d+)", res).group(1)) if re.search(r"分数[：:]\s*(\d+)", res) else 0
                    st.session_state.score_records.append({"学生": s_name, "分数": score, "评价": res})

def page_analysis():
    nav_bar(show_back=True)
    st.subheader("📈 成绩与分析报告")
    if not st.session_state.score_records:
        st.warning("当前无批改记录")
        return
    st.dataframe(st.session_state.score_records, use_container_width=True)
    scores = [r["分数"] for r in st.session_state.score_records]
    col1, col2 = st.columns(2)
    with col1:
        st.metric("平均分", f"{np.mean(scores):.1f}")
        fig, ax = plt.subplots(figsize=(8, 4))
        ax.hist(scores, bins=range(0, 110, 10), color='#4F8BF9', edgecolor='white')
        st.pyplot(fig)
    with col2:
        st.download_button("导出成绩记录 (CSV)", str(st.session_state.score_records), "scores.csv")

# --- 7. 路由逻辑 ---
route = {
    "首页": page_home, "大纲": page_syllabus, "日历": page_calendar, 
    "方案": page_program, "批卷": page_grading, "分析": page_analysis
}
current = st.query_params.get("page", "首页")
route.get(current, page_home)()