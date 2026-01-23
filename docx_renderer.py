from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re

class MarkdownToDocx:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        # Base style config if needed
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'SimSun' # Default to SimSun for general compatibility
        font.size = Pt(12)
        # Setup element defaults for Chinese font rendering
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    def add_markdown_content(self, markdown_text):
        if not markdown_text: return

        # 0. Pre-process cleanup
        # Replace <br> with a placeholder to prevent splitting table rows
        # We use a rare string sequence that won't appear in normal text
        text = re.sub(r'<br\s*/?>', '{{BR}}', markdown_text)
        
        lines = text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue

            # 1. Headers
            if line.startswith('#'):
                level = len(line.split(' ')[0])
                text = line.lstrip('#').strip()
                # Limit to Heading 1-3 to avoid errors if 4+ not defined
                if level > 3: level = 3 
                self.doc.add_heading(text, level=level)
                i += 1
                continue

            # 2. Tables
            # Detection logic: A line containing pipes, followed by a separator line
            # OR a line that looks like a row if we are already in "table mode" (heuristic)
            if '|' in line:
                # Look ahead to see if it's a valid table (needs to be followed by divider or be part of one)
                table_lines = []
                # Check for divider in next line or previous lines if we missed start
                # Simplify: valid block is a group of lines containing '|'
                
                # Consume all consecutive lines containing '|'
                has_separator = False
                while i < len(lines) and ('|' in lines[i].strip()):
                    l = lines[i].strip()
                    if '---' in l: has_separator = True
                    table_lines.append(l)
                    i += 1
                
                # If we found a separator, or it looks strongly like a table (multiple columns)
                # But to avoid false positives with normal text containing |, we require at least 2 lines or a separator
                if (len(table_lines) >= 2 and any('|' in l for l in table_lines)) or has_separator:
                     self._parse_table(table_lines)
                else:
                    # Fallback to normal text for single line without separator
                    for tl in table_lines:
                         p = self.doc.add_paragraph()
                         self._add_run_with_formatting(p, tl)
                continue


            # 3. List Items
            # Bullet points
            if re.match(r'^[\-\*]\s+', line):
                clean_text = re.sub(r'^[\-\*]\s+', '', line)
                p = self.doc.add_paragraph(style='List Bullet')
                self._add_run_with_formatting(p, clean_text)
                i += 1
                continue
                
            # 4. Numbered Lists (Relaxed regex: allow zero spaces after dot)
            if re.match(r'^\d+\.\s*', line):
                clean_text = re.sub(r'^\d+\.\s*', '', line)
                p = self.doc.add_paragraph(style='List Number')
                self._add_run_with_formatting(p, clean_text)
                i += 1
                continue

            # 5. Normal Text (with bold support)
            p = self.doc.add_paragraph()
            self._add_run_with_formatting(p, line)
            i += 1

    def _parse_table(self, table_lines):
        # Basic Markdown Table Parser
        # Filter out the separator line |---|---|
        data_rows = [line for line in table_lines if '---' not in line]
        
        if not data_rows: return

        # Normalize rows: remove leading/trailing pipes if exist, then split
        # This handles:
        # | A | B | -> ['A', 'B']
        # A | B     -> ['A', 'B']
        parsed_rows = []
        for line in data_rows:
            # Smart split: strictly by pipe
            cells = line.strip().split('|')
            # Remove empty first/last if result of outer pipes
            if line.strip().startswith('|') and cells: cells.pop(0)
            if line.strip().endswith('|') and cells: cells.pop(-1)
            parsed_rows.append([c.strip() for c in cells])

        if not parsed_rows: return

        max_cols = max(len(r) for r in parsed_rows)
        if max_cols == 0: return 
        
        table = self.doc.add_table(rows=len(parsed_rows), cols=max_cols)
        table.style = 'Table Grid'
        
        for r_idx, row_data in enumerate(parsed_rows):
            row_cells = table.rows[r_idx].cells
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < len(row_cells):
                    cell = row_cells[c_idx]
                    cell._element.clear_content()
                    p = cell.add_paragraph()
                    self._add_run_with_formatting(p, cell_text)


    def _add_run_with_formatting(self, paragraph, text):
        # Improved formatting parser supporting **bold** and *italic*
        # We use a tokenizing approach for robustness
        
        # 1. Clean HTML
        text = re.sub(r'<[^>]+>', '', text)
        
        # 2. Tokenize by formatting markers
        # Split by **...** or *...* or __...__
        # This simple regex matches bold (** or __) and italic (*)
        # Note: non-greedy match
        tokens = re.split(r'(\*\*.*?\*\*|__.*?__|(?<!\*)\*(?!\*).*?(?<!\*)\*(?!\*))', text)
        
        for token in tokens:
            if not token: continue
            
            is_bold = False
            is_italic = False
            content = token
            
            if token.startswith('**') and token.endswith('**'):
                is_bold = True
                content = token[2:-2]
            elif token.startswith('__') and token.endswith('__'):
                is_bold = True
                content = token[2:-2]
            elif token.startswith('*') and token.endswith('*') and len(token) > 2:
                is_italic = True
                content = token[1:-1]
                
            # Handle BR placeholders: split by {{BR}} and add breaks
            sub_parts = content.split('{{BR}}')
            for idx, sub in enumerate(sub_parts):
                if idx > 0: 
                    paragraph.add_run('\n') # Newline in docx run
                run = paragraph.add_run(sub)
                run.bold = is_bold
                run.italic = is_italic

    
    def save(self):
        bio = io.BytesIO()
        self.doc.save(bio)
        return bio.getvalue()

def create_rich_docx(text):
    if not text:
        text = "内容为空"
    converter = MarkdownToDocx()
    converter.add_markdown_content(text)
    return converter.save()
