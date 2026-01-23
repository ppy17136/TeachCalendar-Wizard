import fitz  # PyMuPDF
from docx import Document
import mammoth
import io
import pdfplumber

def extract_text_from_file(file):
    """支持多格式文本提取"""
    try:
        if file.name.endswith(".docx"):
            return "\n".join([p.text for p in Document(file).paragraphs])
        elif file.name.endswith(".pdf"):
            with pdfplumber.open(file) as pdf:
                return "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif file.name.endswith(".doc"):
            return mammoth.convert_to_text(file).value
        return "格式暂不支持"
    except Exception as e:
        return f"解析失败: {str(e)}"


def safe_extract_text(file, max_chars=15000):
    if not file: return ""
    try:
        text_list = []
        if file.name.endswith(".pdf"):
            with fitz.open(stream=file.read(), filetype="pdf") as doc:
                for page in doc:
                    text_list.append(page.get_text())
                    if sum(len(t) for t in text_list) > max_chars: break
            return "".join(text_list)[:max_chars]
            
        elif file.name.endswith(".docx"):
            doc = Document(file)
            for p in doc.paragraphs:
                if p.text.strip(): text_list.append(p.text)
            
            for table in doc.tables:
                for row in table.rows:
                    processed_cells = []
                    for cell in row.cells:
                        content = cell.text
                        # --- 核心改进：非互斥全量替换，涵盖更多 Word 特殊符号 ---
                        # 识别“已选中”符号
                        checked_chars = ['☑', 'þ', '\xfe', '\uf0fe', '☒', '√']
                        # 识别“未选中”符号
                        unchecked_chars = ['☐', '¨', '\xa8', '\uf0a1', '□']
                        
                        for c in checked_chars:
                            content = content.replace(c, '[已选中]')
                        for u in unchecked_chars:
                            content = content.replace(u, '[未选中]')
                        
                        processed_cells.append(content.strip())
                    
                    row_text = [c for c in processed_cells if c]
                    if row_text: text_list.append(" | ".join(row_text))
            
            return "\n".join(text_list)[:max_chars]
        elif file.name.endswith(".doc"):
            return mammoth.convert_to_text(file).value[:max_chars]            
        return ""

    except Exception as e:
        return "" # Suppress error here, let caller handle or just return empty
