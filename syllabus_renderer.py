import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
import io
import re

class SyllabusRenderer:
    """
    Renders structured Syllabus JSON into:
    1. Markdown (for Streamlit UI)
    2. Docx (for Download)
    """
    
    @staticmethod
    def to_markdown(data):
        """Converts structured JSON to Markdown string"""
        if not data: return ""
        
        md = []
        # Title
        if data.get("course_name"):
            md.append(f"# {data['course_name']} 教学大纲")
        
        # 1. Base Info
        md.append("## 一、课程基本信息")
        bi = data.get("base_info", {})
        md.append(f"- **课程名称**: {bi.get('name', '')}")
        md.append(f"- **课程代码**: {bi.get('code', '待定')}")
        md.append(f"- **学分/学时**: {bi.get('credits', '')} / {bi.get('hours', '')}")
        md.append(f"- **课程性质**: {bi.get('type', '')}")
        
        # 2. Objectives
        md.append("\n## 二、课程目标")
        for idx, obj in enumerate(data.get("objectives", []), 1):
            md.append(f"{idx}. {obj}")
            
        # 3. Graduation Requirements Support
        md.append("\n## 三、对毕业要求的支撑")
        md.append("| 毕业要求 | 指标点 | 支撑强度 |")
        md.append("|---|---|---|")
        for item in data.get("grad_support", []):
            md.append(f"| {item.get('req', '')} | {item.get('point', '')} | {item.get('strength', '')} |")
            
        # 4. Content
        md.append("\n## 四、课程内容与学时分配")
        md.append("| 章节 | 内容 | 讲课 | 实验 | 目标 |")
        md.append("|---|---|---|---|---|")
        for ch in data.get("content", []):
            md.append(f"| {ch.get('chapter', '')} | {ch.get('details', '')} | {ch.get('lec_hrs', 0)} | {ch.get('lab_hrs', 0)} | {ch.get('obj_ref', '')} |")
            
        # 5. Assessment
        md.append("\n## 五、考核方式")
        md.append(data.get("assessment", ""))
        
        # 6. Textbooks
        md.append("\n## 六、教材与参考书")
        md.append(data.get("textbook", ""))
        
        return "\n".join(md)

    @staticmethod
    def to_docx(data):
        """Converts structured JSON to Docx bytes"""
        doc = Document()
        
        # Styles Setup
        style = doc.styles['Normal']
        style.font.name = 'SimSun'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        style.font.size = Pt(12)
        
        # Title
        if data.get("course_name"):
            h = doc.add_heading(f"{data['course_name']} 教学大纲", level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # 1. Base Info Table
        doc.add_heading("一、课程基本信息", level=2)
        bi = data.get("base_info", {})
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        
        # Row 1
        table.cell(0, 0).text = "课程名称"
        table.cell(0, 1).text = bi.get("name", "")
        table.cell(0, 2).text = "课程代码"
        table.cell(0, 3).text = bi.get("code", "")
        
        # Row 2
        table.cell(1, 0).text = "适用专业"
        table.cell(1, 1).text = bi.get("major", "")
        table.cell(1, 2).text = "课程性质"
        table.cell(1, 3).text = bi.get("type", "")
        
        # Row 3
        table.cell(2, 0).text = "总学分"
        table.cell(2, 1).text = str(bi.get("credits", ""))
        table.cell(2, 2).text = "总学时"
        table.cell(2, 3).text = str(bi.get("hours", ""))
        
        # Row 4 (Merge for simple layout)
        table.cell(3, 0).text = "先修课程"
        table.cell(3, 1).merge(table.cell(3, 3))
        table.cell(3, 1).text = bi.get("prerequisites", "无")

        # 2. Objectives
        doc.add_heading("二、课程目标", level=2)
        for idx, obj in enumerate(data.get("objectives", []), 1):
            doc.add_paragraph(f"目标{idx}：{obj}")

        # 3. Graduation Requirements Support (Matrix)
        doc.add_heading("三、课程目标对毕业要求的支撑", level=2)
        supports = data.get("grad_support", [])
        if supports:
            t_grad = doc.add_table(rows=1, cols=3)
            t_grad.style = 'Table Grid'
            hdr = t_grad.rows[0].cells
            hdr[0].text = "毕业要求"
            hdr[1].text = "指标点"
            hdr[2].text = "支撑强度/权重"
            
            for item in supports:
                row = t_grad.add_row().cells
                row[0].text = item.get("req", "")
                row[1].text = item.get("point", "")
                row[2].text = item.get("strength", "")
        else:
            doc.add_paragraph("未识别到具体的支撑关系，请根据培养方案补充。")

        # 4. Content (Table)
        doc.add_heading("四、课程内容与学时分配", level=2)
        contents = data.get("content", [])
        if contents:
            t_cont = doc.add_table(rows=1, cols=5)
            t_cont.style = 'Table Grid'
            hdr = t_cont.rows[0].cells
            hdr[0].text = "章节/模块"
            hdr[1].text = "教学内容与要求"
            hdr[2].text = "讲课"
            hdr[3].text = "实验"
            hdr[4].text = "对应目标"
            
            for ch in contents:
                row = t_cont.add_row().cells
                row[0].text = ch.get("chapter", "")
                row[1].text = ch.get("details", "")
                row[2].text = str(ch.get("lec_hrs", ""))
                row[3].text = str(ch.get("lab_hrs", ""))
                row[4].text = ch.get("obj_ref", "")

        # 5. Assessment
        doc.add_heading("五、考核方式与成绩评定", level=2)
        doc.add_paragraph(data.get("assessment", "见正文"))

        # 6. Textbooks
        doc.add_heading("六、教材与参考资料", level=2)
        doc.add_paragraph(data.get("textbook", ""))

        # Save
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()
